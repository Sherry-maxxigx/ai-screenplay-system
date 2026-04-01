from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "交接文档_项目移交说明.md"
OUTPUT = ROOT / "AI剧本系统项目移交说明书.docx"


def set_document_style(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11)


def add_code_paragraph(document: Document, lines: list[str]) -> None:
    for line in lines:
        p = document.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)


def build_doc() -> None:
    document = Document()
    set_document_style(document)

    section = document.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(64)
    section.right_margin = Pt(64)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()

    in_code_block = False
    code_lines: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code_block:
                add_code_paragraph(document, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not line.strip():
            document.add_paragraph("")
            continue

        if line.startswith("# "):
            document.add_paragraph(line[2:].strip(), style="Title")
            continue

        if line.startswith("## "):
            document.add_paragraph(line[3:].strip(), style="Heading 1")
            continue

        if line.startswith("### "):
            document.add_paragraph(line[4:].strip(), style="Heading 2")
            continue

        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        document.add_paragraph(line)

    if code_lines:
        add_code_paragraph(document, code_lines)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
